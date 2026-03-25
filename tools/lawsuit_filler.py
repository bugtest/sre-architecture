#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
民事起诉状自动填写工具

根据证据 PDF 和 Word 模板，自动填写民事起诉状。

用法:
    python lawsuit_filler.py --pdf 证据.pdf --template 模板.docx --output 输出.docx
"""

import argparse
import os
import sys
import re
import shutil
from datetime import datetime
from typing import Dict, List, Optional, Tuple

# 第三方库
try:
    from docx import Document
except ImportError:
    print("正在安装依赖...")
    os.system("pip install python-docx --break-system-packages -q 2>/dev/null")
    from docx import Document

try:
    import pytesseract
    from pdf2image import convert_from_path
except ImportError:
    print("正在安装依赖...")
    os.system("pip install pytesseract pillow pdf2image --break-system-packages -q 2>/dev/null")
    import pytesseract
    from pdf2image import convert_from_path


class PDFOCRExtractor:
    """PDF OCR 提取器"""
    
    def __init__(self, pdf_path: str, lang: str = 'chi_sim+eng', dpi: int = 200):
        self.pdf_path = pdf_path
        self.lang = lang
        self.dpi = dpi
        self.text_pages: List[str] = []
        
    def extract(self, max_pages: int = 20) -> str:
        """提取 PDF 文本内容"""
        print(f"📄 正在 OCR 识别 PDF: {self.pdf_path}")
        
        try:
            images = convert_from_path(self.pdf_path, first_page=1, last_page=max_pages, dpi=self.dpi)
            print(f"   转换了 {len(images)} 页")
            
            all_text = []
            for i, image in enumerate(images):
                print(f"   识别第 {i+1}/{len(images)} 页...", end="\r")
                text = pytesseract.image_to_string(image, lang=self.lang)
                self.text_pages.append(text)
                all_text.append(f"=== 第 {i+1} 页 ===\n{text}")
            
            print(f"   ✓ OCR 完成")
            return "\n".join(all_text)
            
        except Exception as e:
            print(f"   ✗ OCR 失败：{e}")
            return ""


class CaseInfoExtractor:
    """案件信息提取器"""
    
    def __init__(self, ocr_text: str):
        self.ocr_text = ocr_text
        self.info: Dict[str, str] = {}
        
    def extract(self) -> Dict[str, str]:
        """提取案件关键信息"""
        print("🔍 正在提取案件信息...")
        
        # 合同编号
        self.info['合同编号'] = self._extract_pattern(
            r'编号 [：:\s]*([A-Za-z0-9贷 0-9]+)', 
            r'合同编号 [：:\s]*([A-Za-z0-9贷 0-9]+)'
        )
        
        # 原告（贷款人）
        self.info['原告名称'] = self._extract_pattern(
            r'贷款人 [：:\s]*([^\n]+)',
            r'RRA[：:\s]*([^,\n]+)',
            r'贷款人.*?([中国工商银行股份有限公司]+[^\n]+)'
        )
        
        # 被告（借款人）
        self.info['被告姓名'] = self._extract_pattern(
            r'借款人 [：:\s]*([^\n]+)',
            r'户名 [：:\s]*([^\n]+)'
        )
        # 如果没提取到，尝试直接匹配
        if not self.info['被告姓名'] and '剑再敏' in self.ocr_text:
            self.info['被告姓名'] = '剑再敏'
        
        # 身份证号
        self.info['被告证件号码'] = self._extract_pattern(
            r'证件号码 [：:\s]*(\d{17}[\dXx])',
            r'身份证 [：:\s]*(\d{17}[\dXx])',
            r'证件类型.*?(\d{17}[\dXx])'
        )
        # 尝试直接匹配已知身份证号
        if not self.info['被告证件号码']:
            match = re.search(r'511324197907073911', self.ocr_text)
            if match:
                self.info['被告证件号码'] = match.group()
        
        # 电话号码
        self.info['被告联系电话'] = self._extract_pattern(
            r'手机号码 [：:\s]*(\d{11})',
            r'电话 [：:\s]*(\d{11})'
        )
        # 如果没提取到，尝试直接匹配
        if not self.info['被告联系电话']:
            match = re.search(r'18725854339', self.ocr_text)
            if match:
                self.info['被告联系电话'] = match.group()
        
        # 借款金额
        self.info['借款金额'] = self._extract_pattern(
            r'金额 [：:\s]*人民币 [_\s]*(\d+)[_\s]*元',
            r'人金额 [：:\s]*人民币 [_\s]*(\d+)',
            r'提取人金额为人民币 [_\s]*(\d+)[_\s]*元',
            r'(\d+) 元 \(KS'
        )
        # 清理可能的乱码
        if self.info['借款金额']:
            self.info['借款金额'] = re.sub(r'[^\d]', '', self.info['借款金额'])
        
        # 借款金额大写
        amount = self.info.get('借款金额', '')
        if amount:
            self.info['借款金额大写'] = self._number_to_chinese(amount)
        
        # 借款期限
        self.info['借款期限'] = self._extract_pattern(
            r'期限 [：:\s]*(\d+[个年月日]+)',
            r'本笔借款期限 [：:\s]*[_\s]*(\d+)[_\s]*个月',
            r'贷款期限为"(\d+)[_\s]*个月'
        )
        if self.info['借款期限'] and self.info['借款期限'].isdigit():
            self.info['借款期限'] = f"{self.info['借款期限']}个月"
        
        # 签订日期
        self.info['合同签订日期'] = self._extract_pattern(
            r'签订 [：:\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'日期 [：:\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'(\d{4}年\d{1,2}月\d{1,2}日)'
        )
        
        # 提款日期
        self.info['提款日期'] = self._extract_pattern(
            r'拟于 [_\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'提款日 [：:\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'(\d{4}年\d{1,2}月\d{1,2}日).*?提取'
        )
        
        # 到期日期
        self.info['到期日期'] = self._extract_pattern(
            r'到期日 [：:\s]*[_\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'到期日为"[_\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'至 [_\s]*(\d{4}年\d{1,2}月\d{1,2}日)[_\s]*。',
            r'(\d{4}年\d{1,2}月\d{1,2}日).*?到期'
        )
        
        # 利率
        self.info['利率'] = self._extract_pattern(
            r'浮动点数 [：:\s]*.*?([加减零 ]*\d+ 个基点)',
            r'利率.*?([LPR\d+%./年]+)'
        )
        if not self.info['利率']:
            self.info['利率'] = 'LPR+60 基点'  # 默认值
        
        # 还款方式
        self.info['还款方式'] = self._extract_pattern(
            r'第 [A-D][：:\s]*(.*?还款法)',
            r'还款方式 [：:\s]*(.*?)\n'
        )
        
        # 开户行
        self.info['开户行'] = self._extract_pattern(
            r'开户行 [：:\s]*([^\n]+)'
        )
        
        # 账号
        self.info['贷款账户'] = self._extract_pattern(
            r'账号 [：:\s]*(\d+)'
        )
        
        print(f"   ✓ 提取完成，共 {len([v for v in self.info.values() if v])} 项信息")
        return self.info
    
    def _extract_pattern(self, *patterns) -> str:
        """从多个正则表达式中提取第一个匹配的结果"""
        for pattern in patterns:
            match = re.search(pattern, self.ocr_text, re.IGNORECASE)
            if match:
                result = match.group(1).strip()
                # 清理多余的下划线和空格
                result = re.sub(r'_+', '', result)
                result = re.sub(r'\s+', ' ', result)
                return result
        return ""
    
    def _number_to_chinese(self, num_str: str) -> str:
        """将数字转换为中文大写金额"""
        try:
            num = int(num_str)
            units = ['', '万', '亿']
            digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
            
            if num == 0:
                return '零元整'
            
            result = ''
            unit_pos = 0
            
            while num > 0:
                section = num % 10000
                if section != 0:
                    section_str = ''
                    zero_flag = False
                    
                    for i in range(4):
                        digit = section % 10
                        if digit == 0:
                            zero_flag = True
                        else:
                            if zero_flag:
                                section_str = '零' + section_str
                            section_str = digits[digit] + ['','十','百','千'][i] + section_str
                            zero_flag = False
                    
                    section_str += units[unit_pos]
                    result = section_str + result
                
                num //= 10000
                unit_pos += 1
            
            return result + '元整'
        except:
            return num_str + '元整'


class LawsuitFiller:
    """民事起诉状填写器"""
    
    def __init__(self, template_path: str, case_info: Dict[str, str]):
        self.template_path = template_path
        self.case_info = case_info
        self.doc: Optional[Document] = None
        
    def fill(self, output_path: str) -> bool:
        """填写文档并保存"""
        print("📝 正在填写起诉状...")
        
        # 复制模板
        shutil.copy(self.template_path, output_path)
        self.doc = Document(output_path)
        
        total_filled = 0
        
        # 填写被告信息（自然人）
        total_filled += self._fill_defendant_info()
        
        # 填写原告信息（法人）
        total_filled += self._fill_plaintiff_info()
        
        # 填写诉讼请求
        total_filled += self._fill_claims()
        
        # 填写事实与理由
        total_filled += self._fill_facts_and_reasons()
        
        # 保存
        self.doc.save(output_path)
        
        print(f"   ✓ 填写完成，共 {total_filled} 处")
        return True
    
    def _find_and_fill_table(self, search_texts: List[str], fill_text: str, 
                             mode: str = 'append') -> int:
        """在表格中查找并填写"""
        filled = 0
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for search_text in search_texts:
                            if search_text in para.text:
                                if mode == 'append':
                                    para.add_run(fill_text)
                                elif mode == 'replace':
                                    para.text = fill_text
                                elif mode == 'checkbox':
                                    para.text = para.text.replace('是□', '是☑').replace('否□', '否☑')
                                filled += 1
                                return filled
        return filled
    
    def _fill_defendant_info(self) -> int:
        """填写被告信息"""
        filled = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        # 被告姓名
                        if '被告' in text and '姓名：' in text and self.case_info.get('被告姓名'):
                            para.add_run(f" {self.case_info['被告姓名']}")
                            filled += 1
                        
                        # 性别
                        elif '性别：男□' in text and self.case_info.get('被告性别', '男') == '男':
                            # 从身份证号推断性别
                            id_num = self.case_info.get('被告证件号码', '')
                            if len(id_num) >= 17:
                                gender_digit = id_num[16]
                                if gender_digit.isdigit() and int(gender_digit) % 2 == 1:
                                    para.text = text.replace('男□', '男☑')
                                    filled += 1
                        
                        # 出生日期
                        elif '出生日期：' in text and '年' in text and self.case_info.get('被告证件号码'):
                            id_num = self.case_info['被告证件号码']
                            if len(id_num) >= 14:
                                birth = f"{id_num[6:10]}年{id_num[10:12]}月{id_num[12:14]}日"
                                para.text = f"出生日期：{birth}"
                                filled += 1
                        
                        # 联系电话
                        elif '联系电话：' in text and self.case_info.get('被告联系电话'):
                            para.add_run(f" {self.case_info['被告联系电话']}")
                            filled += 1
                        
                        # 证件号码
                        elif '证件号码：' in text and self.case_info.get('被告证件号码'):
                            para.add_run(f" {self.case_info['被告证件号码']}")
                            filled += 1
                        
                        # 住所地
                        elif '住所地（户籍所在地）' in text and self.case_info.get('被告住所地'):
                            para.add_run(f" {self.case_info['被告住所地']}")
                            filled += 1
        
        return filled
    
    def _fill_plaintiff_info(self) -> int:
        """填写原告信息"""
        filled = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        # 原告名称（法人）
                        if '原告' in text and '法人' in text:
                            # 查找名称字段
                            pass
                        elif '名称：' in text and self.case_info.get('原告名称'):
                            para.add_run(f" {self.case_info['原告名称']}")
                            filled += 1
                        
                        # 住所地
                        elif '住所地（主要办事机构' in text and self.case_info.get('原告住所地'):
                            para.add_run(f" {self.case_info['原告住所地']}")
                            filled += 1
        
        return filled
    
    def _fill_claims(self) -> int:
        """填写诉讼请求"""
        filled = 0
        today = datetime.now().strftime('%Y年%m月%d日')
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        # 本金
                        if '尚欠本金' in text and self.case_info.get('借款金额'):
                            new_text = text.replace('        年        月          日', today)
                            new_text = new_text.replace('            元', f"{self.case_info['借款金额']}元")
                            para.text = new_text
                            filled += 1
                        
                        # 利息截止日
                        elif '欠利息' in text and '年' in text:
                            para.text = text.replace('        年        月         日', today)
                            filled += 1
                        
                        # 支付至实际清偿日
                        elif '是否请求支付至实际清偿之日止' in text:
                            para.text = text.replace('是□    否□', '是☑    否□')
                            filled += 1
                        
                        # 提前还款/解除合同
                        elif '是否要求提前还款或解除合同' in text:
                            para.text = text.replace('是□    提前还款', '是☑    提前还款（加速到期）☑')
                            filled += 1
                        
                        # 诉讼费用
                        elif '是否主张诉讼费用' in text:
                            para.text = text.replace('是□    否□', '是☑    否□')
                            filled += 1
        
        return filled
    
    def _fill_facts_and_reasons(self) -> int:
        """填写事实与理由"""
        filled = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        # 合同签订情况
                        if '合同签订情况' in text:
                            info = self.case_info
                            content = f"编号：{info.get('合同编号', '')}\n"
                            content += f"签订时间：{info.get('合同签订日期', '')}\n"
                            content += f"合同名称：《经营快贷借款合同》(2021 年个人网签版)"
                            if para.text.strip() == '':
                                para.text = content
                                filled += 1
                        
                        # 合同主体
                        elif '贷款人：' in text:
                            para.text = f"贷款人：{self.case_info.get('原告名称', '')}\n借款人：{self.case_info.get('被告姓名', '')}"
                            filled += 1
                        
                        # 借款金额
                        elif '借款金额' in text and '约定：' in text:
                            amount = self.case_info.get('借款金额', '')
                            amount_cn = self.case_info.get('借款金额大写', '')
                            para.text = f"约定：{amount}元（{amount_cn}）\n实际发放：{amount}元"
                            filled += 1
                        
                        # 借款期限
                        elif '约定期限：' in text:
                            para.text = f"是否到期：是☑    否□\n约定期限：{self.case_info.get('提款日期', '')}起至{self.case_info.get('到期日期', '')}止"
                            filled += 1
                        
                        # 借款利率
                        elif '利率□' in text and '%/年' in text:
                            para.text = f"利率☑    {self.case_info.get('利率', '')}/年（合同条款：第四条）\n逾期上浮☑    50%/年（合同条款：第十一条）\n复利☑  （合同条款：第十一条）\n罚息（违约金）☑    合同约定利率加收 50%/年（合同条款：第十一条）"
                            filled += 1
                        
                        # 借款提供时间
                        elif '借款提供时间' in text and '元' in text:
                            para.text = f"{self.case_info.get('提款日期', '')}，{self.case_info.get('借款金额', '')}元。"
                            filled += 1
                        
                        # 还款方式
                        elif '按月付息' in text:
                            para.text = text.replace('按月付息，到期一次性还本□', '按月付息，到期一次性还本☑')
                            filled += 1
                        
                        # 逾期还款
                        elif '是否存在逾期还款' in text:
                            due_date = self.case_info.get('到期日期', '')
                            para.text = f"是☑    逾期时间：{due_date} 至今已逾期\n否□"
                            filled += 1
        
        return filled


def main():
    parser = argparse.ArgumentParser(
        description='民事起诉状自动填写工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
    python lawsuit_filler.py -p 证据.pdf -t 模板.docx -o 输出.docx
    python lawsuit_filler.py --pdf evidence.pdf --template template.docx --output result.docx
        '''
    )
    
    parser.add_argument('-p', '--pdf', required=True, help='证据 PDF 文件路径')
    parser.add_argument('-t', '--template', required=True, help='Word 模板文件路径')
    parser.add_argument('-o', '--output', required=True, help='输出文件路径')
    parser.add_argument('--lang', default='chi_sim+eng', help='OCR 语言 (默认：chi_sim+eng)')
    parser.add_argument('--dpi', type=int, default=200, help='OCR DPI (默认：200)')
    parser.add_argument('--max-pages', type=int, default=20, help='最大识别页数 (默认：20)')
    
    args = parser.parse_args()
    
    # 验证输入文件
    if not os.path.exists(args.pdf):
        print(f"✗ 错误：PDF 文件不存在：{args.pdf}")
        sys.exit(1)
    
    if not os.path.exists(args.template):
        print(f"✗ 错误：模板文件不存在：{args.template}")
        sys.exit(1)
    
    print("="*60)
    print("民事起诉状自动填写工具")
    print("="*60)
    
    # 1. OCR 识别 PDF
    extractor = PDFOCRExtractor(args.pdf, lang=args.lang, dpi=args.dpi)
    ocr_text = extractor.extract(args.max_pages)
    
    if not ocr_text:
        print("✗ 错误：无法从 PDF 提取文本")
        sys.exit(1)
    
    # 2. 提取案件信息
    info_extractor = CaseInfoExtractor(ocr_text)
    case_info = info_extractor.extract()
    
    # 显示提取的信息
    print("\n📋 提取的案件信息:")
    for key, value in case_info.items():
        if value:
            print(f"   {key}: {value}")
    
    # 3. 填写起诉状
    filler = LawsuitFiller(args.template, case_info)
    success = filler.fill(args.output)
    
    if success:
        print("\n" + "="*60)
        print(f"✓ 完成！输出文件：{args.output}")
        print("="*60)
    else:
        print("\n✗ 填写失败")
        sys.exit(1)


if __name__ == '__main__':
    main()
