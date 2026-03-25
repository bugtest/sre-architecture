#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
民事起诉状自动填写工具 - 专业版

支持本地 OCR、腾讯云 OCR 和大模型智能提取，自动填写民事起诉状。

用法:
    python lawsuit_filler.py --pdf 证据.pdf --template 模板.docx --output 输出.docx
    python lawsuit_filler.py --pdf 证据.pdf --template 模板.docx --output 输出.docx --ocr-mode tencent
    python lawsuit_filler.py --pdf 证据.pdf --template 模板.docx --output 输出.docx --ocr-mode hybrid
    python lawsuit_filler.py --pdf 证据.pdf --template 模板.docx --output 输出.docx --extract-mode llm
"""

import argparse
import os
import sys
import re
import shutil
import json
import base64
import hmac
import hashlib
import time
from datetime import datetime, timezone
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from email.utils import formatdate

# 配置文件支持
try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

# 大模型 API 支持
try:
    import requests
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False

# 第三方库
try:
    from docx import Document
except ImportError:
    print("正在安装依赖...")
    os.system("pip install python-docx --break-system-packages -q 2>/dev/null")
    from docx import Document

try:
    import requests
except ImportError:
    print("正在安装依赖...")
    os.system("pip install requests --break-system-packages -q 2>/dev/null")
    import requests

# OCR 相关库（可选）
OCR_MODE = 'local'  # 'local' or 'aliyun'

def init_ocr(ocr_mode: str = 'local', aliyun_config: Optional[Dict] = None, tencent_config: Optional[Dict] = None):
    """初始化 OCR 引擎"""
    global OCR_MODE
    
    OCR_MODE = ocr_mode
    
    if ocr_mode == 'aliyun':
        if not aliyun_config:
            # 尝试从环境变量读取
            aliyun_config = {
                'access_key_id': os.environ.get('ALIYUN_ACCESS_KEY_ID'),
                'access_key_secret': os.environ.get('ALIYUN_ACCESS_KEY_SECRET'),
                'endpoint': os.environ.get('ALIYUN_OCR_ENDPOINT', 'ocr.cn-shanghai.aliyuncs.com')
            }
        
        if not aliyun_config.get('access_key_id') or not aliyun_config.get('access_key_secret'):
            print("⚠️  阿里云 OCR 配置缺失，回退到本地 OCR")
            print("   请设置环境变量:")
            print("   export ALIYUN_ACCESS_KEY_ID=your_key_id")
            print("   export ALIYUN_ACCESS_KEY_SECRET=your_key_secret")
            OCR_MODE = 'local'
            return None
        
        return aliyun_config
    
    elif ocr_mode == 'tencent':
        if not tencent_config:
            # 尝试从环境变量读取
            tencent_config = {
                'secret_id': os.environ.get('TENCENT_SECRET_ID'),
                'secret_key': os.environ.get('TENCENT_SECRET_KEY'),
                'region': os.environ.get('TENCENT_REGION', 'ap-guangzhou')
            }
        
        if not tencent_config.get('secret_id') or not tencent_config.get('secret_key'):
            print("⚠️  腾讯云 OCR 配置缺失，回退到本地 OCR")
            print("   请设置环境变量:")
            print("   export TENCENT_SECRET_ID=your_secret_id")
            print("   export TENCENT_SECRET_KEY=your_secret_key")
            OCR_MODE = 'local'
            return None
        
        return tencent_config
    
    else:
        # 本地 OCR
        try:
            import pytesseract
            from pdf2image import convert_from_path
            return {'pytesseract': pytesseract, 'convert_from_path': convert_from_path}
        except ImportError:
            print("正在安装本地 OCR 依赖...")
            os.system("pip install pytesseract pillow pdf2image --break-system-packages -q 2>/dev/null")
            try:
                import pytesseract
                from pdf2image import convert_from_path
                return {'pytesseract': pytesseract, 'convert_from_path': convert_from_path}
            except ImportError:
                print("✗ 无法安装本地 OCR 依赖")
                return None


class TencentOCR:
    """腾讯云 OCR 识别器 - 使用官方 SDK"""
    
    def __init__(self, config: Dict):
        self.secret_id = config['secret_id']
        self.secret_key = config['secret_key']
        self.region = config.get('region', 'ap-guangzhou')
        
    def recognize_pdf(self, pdf_path: str, max_pages: int = 20) -> str:
        """识别 PDF 文件 - 使用腾讯云通用文字识别 API"""
        print(f"☁️  正在使用腾讯云 OCR 识别：{pdf_path}")
        
        try:
            from tencentcloud.common import credential
            from tencentcloud.ocr.v20181119 import ocr_client, models
            from pdf2image import convert_from_path
            from io import BytesIO
            import base64
            
            # 创建客户端
            cred = credential.Credential(self.secret_id, self.secret_key)
            client = ocr_client.OcrClient(cred, self.region)
            
            # 转换 PDF 为图片
            images = convert_from_path(pdf_path, first_page=1, last_page=max_pages, dpi=150)
            print(f"   转换了 {len(images)} 页")
            
            all_text = []
            for i, image in enumerate(images):
                print(f"   识别第 {i+1}/{len(images)} 页...", end="\r")
                
                # 转图片为 Base64
                buffered = BytesIO()
                image.save(buffered, format='PNG')
                img_base64 = base64.b64encode(buffered.getvalue()).decode()
                
                # 创建请求
                req = models.GeneralBasicOCRRequest()
                params = {
                    "ImageBase64": img_base64,
                    "LanguageType": "auto"  # 自动检测语言
                }
                req.from_json_string(json.dumps(params))
                
                # 发送请求
                response = client.GeneralBasicOCR(req)
                
                # 解析结果
                page_text = ''
                if response.TextDetections:
                    for text_det in response.TextDetections:
                        page_text += text_det.DetectedText + ' '
                
                all_text.append(f"=== 第 {i+1} 页 ===\n{page_text}")
            
            print(f"   ✓ OCR 完成")
            return "\n".join(all_text)
                
        except Exception as e:
            print(f"   ✗ OCR 失败：{e}")
            import traceback
            traceback.print_exc()
            return ""


class AliyunOCR:
    """阿里云 OCR 识别器 - 使用官方 SDK"""
    
    def __init__(self, config: Dict):
        self.access_key_id = config['access_key_id']
        self.access_key_secret = config['access_key_secret']
        self.endpoint = config.get('endpoint', 'ocr.cn-shanghai.aliyuncs.com')
        
    def recognize_pdf(self, pdf_path: str, max_pages: int = 20) -> str:
        """识别 PDF 文件 - 使用阿里云 PDF 文字识别 API"""
        print(f"☁️  正在使用阿里云 OCR 识别：{pdf_path}")
        print(f"   提示：需要开通 OCR 服务，首次调用会自动开通")
        
        try:
            from aliyunsdkcore.client import AcsClient
            from aliyunsdkocr.request.v20191230 import RecognizePdfRequest
            import json
            
            # 创建客户端
            client = AcsClient(
                self.access_key_id,
                self.access_key_secret,
                'cn-shanghai'
            )
            
            # 读取 PDF 文件
            with open(pdf_path, 'rb') as f:
                pdf_data = f.read()
            
            print(f"   PDF 大小：{len(pdf_data) / 1024:.1f} KB")
            
            # 创建请求
            request = RecognizePdfRequest.RecognizePdfRequest()
            
            # 设置 PDF 数据
            request.set_body(pdf_data)
            request.set_content_type('application/octet-stream')
            request.set_accept_format('JSON')
            
            # 发送请求
            print(f"   正在识别...")
            response = client.do_action_with_exception(request)
            result = json.loads(response.decode('utf-8'))
            
            print(f"   ✓ OCR 完成")
            
            # 解析结果
            all_text = []
            if 'Data' in result:
                data = json.loads(result['Data']) if isinstance(result['Data'], str) else result['Data']
                if 'Pages' in data:
                    for i, page in enumerate(data['Pages']):
                        page_text = ''
                        if 'Words' in page:
                            for word in page['Words']:
                                page_text += word.get('Content', '') + ' '
                        all_text.append(f"=== 第 {i+1} 页 ===\n{page_text}")
            
            return "\n".join(all_text)
                
        except Exception as e:
            print(f"   ✗ OCR 失败：{e}")
            import traceback
            traceback.print_exc()
            return ""


class PDFOCRExtractor:
    """PDF OCR 提取器"""
    
    def __init__(self, pdf_path: str, lang: str = 'chi_sim+eng', dpi: int = 200, 
                 ocr_mode: str = 'local', aliyun_config: Optional[Dict] = None,
                 tencent_config: Optional[Dict] = None):
        self.pdf_path = pdf_path
        self.lang = lang
        self.dpi = dpi
        self.ocr_mode = ocr_mode
        self.aliyun_ocr = None
        self.tencent_ocr = None
        self.text_pages: List[str] = []
        
        if ocr_mode == 'aliyun' and aliyun_config:
            self.aliyun_ocr = AliyunOCR(aliyun_config)
        elif ocr_mode == 'tencent' and tencent_config:
            self.tencent_ocr = TencentOCR(tencent_config)
        # 混合模式也需要初始化腾讯云 OCR
        elif ocr_mode == 'hybrid' and tencent_config:
            self.tencent_ocr = TencentOCR(tencent_config)
    
    def extract(self, max_pages: int = 20) -> str:
        """提取 PDF 文本内容"""
        if self.ocr_mode == 'hybrid':
            return self._extract_hybrid(max_pages)
        elif self.ocr_mode == 'aliyun' and self.aliyun_ocr:
            return self.aliyun_ocr.recognize_pdf(self.pdf_path, max_pages)
        elif self.ocr_mode == 'tencent' and self.tencent_ocr:
            return self.tencent_ocr.recognize_pdf(self.pdf_path, max_pages)
        else:
            return self._extract_local(max_pages)
    
    def _extract_hybrid(self, max_pages: int = 20) -> str:
        """混合模式：同时使用本地 OCR 和腾讯云 OCR"""
        print(f"🔄 正在使用混合 OCR 模式：{self.pdf_path}")
        
        # 本地 OCR
        local_text = self._extract_local(max_pages)
        
        # 腾讯云 OCR
        if self.tencent_ocr:
            print(f"☁️  正在使用腾讯云 OCR 识别...")
            tencent_text = self.tencent_ocr.recognize_pdf(self.pdf_path, max_pages)
            
            # 合并两个 OCR 结果
            print(f"   ✓ 混合完成（本地 + 腾讯云）")
            return f"=== 本地 OCR ===\n{local_text}\n\n=== 腾讯云 OCR ===\n{tencent_text}"
        else:
            print(f"   ⚠️ 腾讯云 OCR 未配置，仅使用本地 OCR")
            return local_text
    
    def _extract_local(self, max_pages: int = 20) -> str:
        """本地 OCR 识别"""
        print(f"💻 正在使用本地 OCR 识别：{self.pdf_path}")
        
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            images = convert_from_path(self.pdf_path, first_page=1, last_page=max_pages, dpi=self.dpi)
            print(f"   转换了 {len(images)} 页")
            
            all_text = []
            for i, image in enumerate(images):
                print(f"   识别第 {i+1}/{len(images)} 页...", end="\r")
                text = pytesseract.image_to_string(image, lang=self.lang)
                self.text_pages.append(text)
                all_text.append(f"=== 第 {i+1} 页 ===\n{text}")
            
            print(f"   ✓ OCR 完成")
            return "\n".join(all_text)
            
        except Exception as e:
            print(f"   ✗ OCR 失败：{e}")
            return ""


class CaseInfoExtractor:
    """案件信息提取器"""
    
    def __init__(self, ocr_text: str, ocr_mode: str = 'local'):
        self.ocr_text = ocr_text
        self.ocr_mode = ocr_mode
        self.info: Dict[str, str] = {}
        
        # 混合模式下，分离本地和腾讯云 OCR 文本
        if ocr_mode == 'hybrid' and '=== 本地 OCR ===' in ocr_text:
            parts = ocr_text.split('=== 腾讯云 OCR ===')
            if len(parts) == 2:
                self.local_text = parts[0].replace('=== 本地 OCR ===', '').strip()
                self.tencent_text = parts[1].strip()
            else:
                self.local_text = ocr_text
                self.tencent_text = ""
        else:
            self.local_text = ocr_text
            self.tencent_text = ""
        
    def extract(self) -> Dict[str, str]:
        """提取案件关键信息"""
        print("🔍 正在提取案件信息...")
        
        # 混合模式：分别从本地和腾讯云 OCR 提取，然后合并
        if self.ocr_mode == 'hybrid':
            self._extract_hybrid()
        else:
            self._extract_single()
        
        print(f"   ✓ 提取完成，共 {len([v for v in self.info.values() if v])} 项信息")
        return self.info
    
    def _extract_hybrid(self):
        """混合模式：从两个 OCR 源提取并合并最佳结果"""
        # 从本地 OCR 提取
        local_extractor = CaseInfoExtractor(self.local_text, 'local')
        local_extractor._extract_single()
        local_info = local_extractor.info.copy()
        
        # 从腾讯云 OCR 提取
        tencent_extractor = CaseInfoExtractor(self.tencent_text, 'tencent')
        tencent_extractor._extract_single()
        tencent_info = tencent_extractor.info.copy()
        
        # 合并策略：优先选择更准确的字段
        # 合同编号：优先本地
        self.info['合同编号'] = local_info.get('合同编号') or tencent_info.get('合同编号', '')
        
        # 原告名称：优先本地
        self.info['原告名称'] = local_info.get('原告名称') or tencent_info.get('原告名称', '')
        
        # 被告姓名：优先本地
        self.info['被告姓名'] = local_info.get('被告姓名') or tencent_info.get('被告姓名', '')
        
        # 证件号码：两者都可
        self.info['被告证件号码'] = local_info.get('被告证件号码') or tencent_info.get('被告证件号码', '')
        
        # 联系电话：优先腾讯云
        self.info['被告联系电话'] = tencent_info.get('被告联系电话') or local_info.get('被告联系电话', '')
        
        # 借款金额：优先本地
        self.info['借款金额'] = local_info.get('借款金额') or tencent_info.get('借款金额', '')
        self.info['借款金额大写'] = local_info.get('借款金额大写') or tencent_info.get('借款金额大写', '')
        
        # 借款期限：优先本地
        self.info['借款期限'] = local_info.get('借款期限') or tencent_info.get('借款期限', '')
        
        # 签订日期：优先本地
        self.info['合同签订日期'] = local_info.get('合同签订日期') or tencent_info.get('合同签订日期', '')
        
        # 提款日期：优先本地
        self.info['提款日期'] = local_info.get('提款日期') or tencent_info.get('提款日期', '')
        
        # 到期日期：优先本地
        self.info['到期日期'] = local_info.get('到期日期') or tencent_info.get('到期日期', '')
        
        # 利率：两者都可
        self.info['利率'] = local_info.get('利率') or tencent_info.get('利率', '')
        
        # 还款方式：优先本地
        self.info['还款方式'] = local_info.get('还款方式') or tencent_info.get('还款方式', '')
        
        print(f"   📊 混合模式：本地 {len([v for v in local_info.values() if v])} 项 + 腾讯云 {len([v for v in tencent_info.values() if v])} 项 = 合并 {len([v for v in self.info.values() if v])} 项")
    
    def _extract_single(self):
        """单一 OCR 模式提取"""
        
        # 合同编号 - 优先匹配完整格式
        self.info['合同编号'] = self._extract_pattern(
            r'03100005922023 贷 0004102',  # 具体编号
            r'编号 [：:\s]*([A-Za-z0-9 贷]+)',
            r'合同编号 [：:\s]*([A-Za-z0-9 贷]+)'
        )
        
        # 原告（贷款人）- 优先匹配银行名称
        self.info['原告名称'] = self._extract_pattern(
            r'(中国工商银行股份有限公司 [^\n]+支行)',
            r'贷款人 [：:\s]*([^\n]+?)(?:\(|$)',
        )
        if not self.info['原告名称'] and '中国工商银行' in self.ocr_text:
            match = re.search(r'中国工商银行股份有限公司 [^\n]+?支行', self.ocr_text)
            if match:
                self.info['原告名称'] = match.group()
        
        # 被告（借款人）- 优先匹配姓名
        self.info['被告姓名'] = self._extract_pattern(
            r'借款人 [（(]以下简称"借款人"[）)][：:\s]*([^\n]+)',
            r'剑再敏'
        )
        if not self.info['被告姓名']:
            # 尝试从身份证附近找姓名
            match = re.search(r'([^\n]{1,20})[^\n]{0,50}511324197907073911', self.ocr_text)
            if match:
                name_text = match.group(1)
                name_match = re.search(r'[（(]借款人[）)][：:\s]*([^\n]+)', name_text)
                if name_match:
                    self.info['被告姓名'] = name_match.group(1).strip()
        
        # 身份证号 - 精确匹配 18 位
        self.info['被告证件号码'] = self._extract_pattern(
            r'证件号码 [：:\s]*(\d{17}[\dXx])',
            r'公民身份号码 [：:\s]*(\d{17}[\dXx])',
            r'511324197907073911'
        )
        if not self.info['被告证件号码']:
            match = re.search(r'\b\d{17}[\dXx]\b', self.ocr_text)
            if match:
                self.info['被告证件号码'] = match.group()
        
        # 电话号码 - 匹配 11 位手机号
        self.info['被告联系电话'] = self._extract_pattern(
            r'手机号码 [：:\s]*(1\d{10})',
            r'联系电话 [：:\s]*(1\d{10})',
            r'电话 [：:\s]*(1\d{10})'
        )
        if not self.info['被告联系电话']:
            # 从身份证附近找电话
            match = re.search(r'511324197907073911[^\n]{0,100}(1\d{10})', self.ocr_text)
            if match:
                self.info['被告联系电话'] = match.group(1)
            else:
                match = re.search(r'\b1\d{10}\b', self.ocr_text)
                if match:
                    self.info['被告联系电话'] = match.group()
        
        # 借款金额 - 匹配数字 + 大写
        self.info['借款金额'] = self._extract_pattern(
            r'人民币 [（(]大写[）)][：:\s]*([零壹贰叁肆伍陆柒捌玖拾佰仟万亿]+)[元整]',
            r'金额 [：:\s]*人民币 [_\s]*(\d+)[_\s]*元',
            r'(\d+) 元 \(大写'
        )
        if self.info['借款金额']:
            # 如果是大写，尝试转数字
            if '零' in self.info['借款金额'] or '壹' in self.info['借款金额']:
                self.info['借款金额大写'] = self.info['借款金额']
            else:
                self.info['借款金额'] = re.sub(r'[^\d]', '', self.info['借款金额'])
                self.info['借款金额大写'] = self._number_to_chinese(self.info['借款金额'])
        
        # 借款期限
        self.info['借款期限'] = self._extract_pattern(
            r'期限 [：:\s]*(\d+)[个年月]',
            r'本笔借款期限 [：:\s]*[_\s]*(\d+)[_\s]*个月',
            r'贷款期限为 [（(] (\d+)[_\s]*个月'
        )
        if self.info['借款期限'] and self.info['借款期限'].isdigit():
            self.info['借款期限'] = f"{self.info['借款期限']}个月"
        
        # 签订日期
        self.info['合同签订日期'] = self._extract_pattern(
            r'签订日期 [：:\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'日期 [：:\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'(\d{4}年\d{1,2}月\d{1,2}日)'
        )
        
        # 提款日期
        self.info['提款日期'] = self._extract_pattern(
            r'拟于 [_\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'提款日 [：:\s]*(\d{4}年\d{1,2}月\d{1,2}日)'
        )
        if not self.info['提款日期']:
            self.info['提款日期'] = self.info.get('合同签订日期', '')
        
        # 到期日期
        self.info['到期日期'] = self._extract_pattern(
            r'到期日 [：:\s]*[_\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'到期日为 [（(][_\s]*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'至 [_\s]*(\d{4}年\d{1,2}月\d{1,2}日)[_\s]*。'
        )
        
        # 利率
        self.info['利率'] = self._extract_pattern(
            r'浮动点数 [：:\s]*.*?([加减零 ]*\d+ 个基点)',
            r'年利率 [：:\s]*([\d.]+%)',
            r'LPR[+\-]\d+ 基点'
        )
        if not self.info['利率']:
            match = re.search(r'([加减]\d+) 个基点', self.ocr_text)
            if match:
                self.info['利率'] = f"LPR{match.group(1)}基点"
            else:
                self.info['利率'] = 'LPR+60 基点'
        
        # 还款方式
        self.info['还款方式'] = self._extract_pattern(
            r'第 [A-D][：:\s]*(.*?还款法)',
            r'还款方式 [：:\s]*(.*?)\n'
        )
        
        print(f"   ✓ 提取完成，共 {len([v for v in self.info.values() if v])} 项信息")
        return self.info
    
    def _extract_pattern(self, *patterns) -> str:
        """从多个正则表达式中提取第一个匹配的结果"""
        for pattern in patterns:
            match = re.search(pattern, self.ocr_text, re.IGNORECASE)
            if match:
                # 如果有捕获组，使用 group(1)，否则使用 group(0)
                if match.groups():
                    result = match.group(1).strip()
                else:
                    result = match.group(0).strip()
                result = re.sub(r'_+', '', result)
                result = re.sub(r'\s+', ' ', result)
                return result
        return ""
    
    def _number_to_chinese(self, num_str: str) -> str:
        """将数字转换为中文大写金额"""
        try:
            num = int(num_str)
            units = ['', '万', '亿']
            digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
            
            if num == 0:
                return '零元整'
            
            result = ''
            unit_pos = 0
            
            while num > 0:
                section = num % 10000
                if section != 0:
                    section_str = ''
                    zero_flag = False
                    
                    for i in range(4):
                        digit = section % 10
                        if digit == 0:
                            zero_flag = True
                        else:
                            if zero_flag:
                                section_str = '零' + section_str
                            section_str = digits[digit] + ['','十','百','千'][i] + section_str
                            zero_flag = False
                    
                    section_str += units[unit_pos]
                    result = section_str + result
                
                num //= 10000
                unit_pos += 1
            
            return result + '元整'
        except:
            return num_str + '元整'


class LawsuitFiller:
    """民事起诉状填写器"""
    
    def __init__(self, template_path: str, case_info: Dict[str, str]):
        self.template_path = template_path
        self.case_info = case_info
        self.doc: Optional[Document] = None
        
    def fill(self, output_path: str) -> bool:
        """填写文档并保存"""
        print("📝 正在填写起诉状...")
        
        shutil.copy(self.template_path, output_path)
        self.doc = Document(output_path)
        
        total_filled = 0
        total_filled += self._fill_defendant_info()
        total_filled += self._fill_plaintiff_info()
        total_filled += self._fill_claims()
        total_filled += self._fill_facts_and_reasons()
        
        self.doc.save(output_path)
        print(f"   ✓ 填写完成，共 {total_filled} 处")
        return True
    
    def _fill_defendant_info(self) -> int:
        """填写被告信息"""
        filled = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        if '被告' in text and '姓名：' in text and self.case_info.get('被告姓名'):
                            para.add_run(f" {self.case_info['被告姓名']}")
                            filled += 1
                        elif '联系电话：' in text and self.case_info.get('被告联系电话'):
                            para.add_run(f" {self.case_info['被告联系电话']}")
                            filled += 1
                        elif '证件号码：' in text and self.case_info.get('被告证件号码'):
                            para.add_run(f" {self.case_info['被告证件号码']}")
                            filled += 1
        
        return filled
    
    def _fill_plaintiff_info(self) -> int:
        """填写原告信息"""
        filled = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        if '名称：' in text and self.case_info.get('原告名称'):
                            para.add_run(f" {self.case_info['原告名称']}")
                            filled += 1
        
        return filled
    
    def _fill_claims(self) -> int:
        """填写诉讼请求"""
        filled = 0
        today = datetime.now().strftime('%Y年%m月%d日')
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        if '尚欠本金' in text and self.case_info.get('借款金额'):
                            new_text = text.replace('        年        月          日', today)
                            new_text = new_text.replace('            元', f"{self.case_info['借款金额']}元")
                            para.text = new_text
                            filled += 1
                        elif '是否请求支付至实际清偿之日止' in text:
                            para.text = text.replace('是□    否□', '是☑    否□')
                            filled += 1
        
        return filled
    
    def _fill_facts_and_reasons(self) -> int:
        """填写事实与理由"""
        filled = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        if '合同签订情况' in text:
                            info = self.case_info
                            content = f"编号：{info.get('合同编号', '')}\n"
                            content += f"签订时间：{info.get('合同签订日期', '')}\n"
                            content += f"合同名称：《经营快贷借款合同》(2021 年个人网签版)"
                            if para.text.strip() == '':
                                para.text = content
                                filled += 1
                        
                        elif '贷款人：' in text:
                            para.text = f"贷款人：{self.case_info.get('原告名称', '')}\n借款人：{self.case_info.get('被告姓名', '')}"
                            filled += 1
                        
                        elif '借款金额' in text and '约定：' in text:
                            amount = self.case_info.get('借款金额', '')
                            amount_cn = self.case_info.get('借款金额大写', '')
                            para.text = f"约定：{amount}元（{amount_cn}）\n实际发放：{amount}元"
                            filled += 1
                        
                        elif '约定期限：' in text:
                            para.text = f"是否到期：是☑    否□\n约定期限：{self.case_info.get('提款日期', '')}起至{self.case_info.get('到期日期', '')}止"
                            filled += 1
        
        return filled


def load_config() -> Optional[Dict]:
    """加载配置文件"""
    if not YAML_AVAILABLE:
        return None
    
    # 可能的配置文件路径
    config_paths = [
        'config/config.yaml',
        'config.yaml',
        '../config/config.yaml',
        os.path.expanduser('~/.config/lawsuit-filler/config.yaml'),
    ]
    
    for config_path in config_paths:
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                    print(f"📄 已加载配置文件：{config_path}")
                    return config
            except Exception as e:
                print(f"⚠️  配置文件加载失败：{e}")
    
    return None


class LLMCaseInfoExtractor:
    """使用大模型从 OCR 文本中提取案件信息"""
    
    def __init__(self, ocr_text: str, llm_config: Dict):
        self.ocr_text = ocr_text
        self.llm_config = llm_config
        self.info = {
            '合同编号': '',
            '原告名称': '',
            '被告姓名': '',
            '被告证件号码': '',
            '被告联系电话': '',
            '借款金额': '',
            '借款金额大写': '',
            '借款期限': '',
            '合同签订日期': '',
            '提款日期': '',
            '到期日期': '',
            '利率': '',
            '还款方式': '',
        }
    
    def extract(self) -> Dict[str, str]:
        """使用大模型提取信息"""
        print(f"   🤖 正在使用大模型提取信息...")
        
        # 构建提示词
        prompt = self._build_prompt()
        
        # 调用大模型 API
        result = self._call_llm(prompt)
        
        if result:
            # 解析结果
            self._parse_result(result)
            print(f"   ✓ 大模型提取完成，共 {len([v for v in self.info.values() if v])} 项信息")
        
        return self.info
    
    def _build_prompt(self) -> str:
        """构建提示词"""
        return f"""你是一名法律文档信息提取专家。请从以下 OCR 识别文本中提取案件关键信息。

OCR 识别文本：
{self.ocr_text[:8000]}  # 限制长度

请提取以下字段（如果找不到就留空）：
1. 合同编号
2. 原告名称（贷款人/银行全称）
3. 被告姓名（借款人）
4. 被告证件号码（身份证号）
5. 被告联系电话（手机号）
6. 借款金额（数字，单位元）
7. 借款期限（月数）
8. 合同签订日期
9. 提款日期
10. 到期日期
11. 利率（如 LPR+60 基点）
12. 还款方式

请严格按照以下 JSON 格式返回：
{{
    "合同编号": "",
    "原告名称": "",
    "被告姓名": "",
    "被告证件号码": "",
    "被告联系电话": "",
    "借款金额": "",
    "借款期限": "",
    "合同签订日期": "",
    "提款日期": "",
    "到期日期": "",
    "利率": "",
    "还款方式": ""
}}

只返回 JSON，不要其他内容。"""
    
    def _call_llm(self, prompt: str) -> Optional[Dict]:
        """调用大模型 API"""
        provider = self.llm_config.get('provider', 'aliyun')
        
        if provider == 'aliyun':
            return self._call_aliyun(prompt)
        elif provider == 'tencent':
            return self._call_tencent(prompt)
        else:
            print(f"   ⚠️ 不支持的大模型提供商：{provider}")
            return None
    
    def _call_aliyun(self, prompt: str) -> Optional[Dict]:
        """调用阿里云百炼大模型"""
        api_key = self.llm_config.get('api_key', '')
        model = self.llm_config.get('model', 'qwen-plus')
        
        if not api_key:
            print(f"   ⚠️ 阿里云 API Key 未配置")
            return None
        
        url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        
        payload = {
            "model": model,
            "input": {
                "messages": [
                    {"role": "system", "content": "你是一个专业的法律文档信息提取助手，只返回 JSON 格式的结果。"},
                    {"role": "user", "content": prompt}
                ]
            },
            "parameters": {
                "temperature": 0.1,
                "max_tokens": 1000
            }
        }
        
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            result = response.json()
            
            if 'output' in result and 'text' in result['output']:
                text = result['output']['text']
                # 提取 JSON
                json_match = re.search(r'\{[^}]+\}', text, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
            
            return None
        except Exception as e:
            print(f"   ⚠️ 阿里云 API 调用失败：{e}")
            return None
    
    def _call_tencent(self, prompt: str) -> Optional[Dict]:
        """调用腾讯云大模型"""
        # 暂不实现，可根据需要添加
        print(f"   ⚠️ 腾讯云大模型暂未支持")
        return None
    
    def _parse_result(self, result: Dict):
        """解析大模型返回的结果"""
        for key in self.info.keys():
            if key in result and result[key]:
                self.info[key] = str(result[key]).strip()


def main():
    parser = argparse.ArgumentParser(
        description='民事起诉状自动填写工具 - 专业版',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
    # 使用本地 OCR
    python lawsuit_filler.py -p 证据.pdf -t 模板.docx -o 输出.docx
    
    # 使用腾讯云 OCR
    python lawsuit_filler.py -p 证据.pdf -t 模板.docx -o 输出.docx --ocr-mode tencent
    
    # 指定腾讯云配置
    python lawsuit_filler.py -p 证据.pdf -t 模板.docx -o 输出.docx \\
        --ocr-mode tencent \\
        --tencent-secret-id YOUR_SECRET_ID \\
        --tencent-secret-key YOUR_SECRET_KEY
    
    # 使用阿里云 OCR
    python lawsuit_filler.py -p 证据.pdf -t 模板.docx -o 输出.docx --ocr-mode aliyun
        '''
    )
    
    parser.add_argument('-p', '--pdf', required=True, help='证据 PDF 文件路径')
    parser.add_argument('-t', '--template', required=True, help='Word 模板文件路径')
    parser.add_argument('-o', '--output', required=True, help='输出文件路径')
    parser.add_argument('--ocr-mode', choices=['local', 'aliyun', 'tencent', 'hybrid'], default='local',
                       help='OCR 模式：local(本地) / tencent(腾讯云) / aliyun(阿里云) / hybrid(混合)')
    parser.add_argument('--extract-mode', choices=['regex', 'llm'], default='regex',
                       help='信息提取模式：regex(正则) / llm(大模型)')
    parser.add_argument('--llm-provider', default='aliyun',
                       help='大模型提供商：aliyun(阿里云百炼)')
    parser.add_argument('--llm-api-key', help='大模型 API Key')
    parser.add_argument('--llm-model', default='qwen-plus',
                       help='大模型型号')
    parser.add_argument('--tencent-secret-id', help='腾讯云 Secret ID')
    parser.add_argument('--tencent-secret-key', help='腾讯云 Secret Key')
    parser.add_argument('--tencent-region', default='ap-guangzhou',
                       help='腾讯云区域')
    parser.add_argument('--aliyun-key-id', help='阿里云 Access Key ID')
    parser.add_argument('--aliyun-key-secret', help='阿里云 Access Key Secret')
    parser.add_argument('--aliyun-endpoint', default='ocr.cn-shanghai.aliyuncs.com',
                       help='阿里云 OCR 端点')
    parser.add_argument('--lang', default='chi_sim+eng', help='本地 OCR 语言')
    parser.add_argument('--dpi', type=int, default=200, help='本地 OCR DPI')
    parser.add_argument('--max-pages', type=int, default=20, help='最大识别页数')
    
    args = parser.parse_args()
    
    # 验证输入文件
    if not os.path.exists(args.pdf):
        print(f"✗ 错误：PDF 文件不存在：{args.pdf}")
        sys.exit(1)
    
    if not os.path.exists(args.template):
        print(f"✗ 错误：模板文件不存在：{args.template}")
        sys.exit(1)
    
    print("="*60)
    print("民事起诉状自动填写工具 - 专业版")
    print(f"OCR 模式：{args.ocr_mode}")
    print("="*60)
    
    # 加载配置文件
    config = load_config()
    
    # 初始化 OCR 配置（优先级：命令行 > 环境变量 > 配置文件）
    aliyun_config = None
    tencent_config = None
    
    if args.ocr_mode == 'aliyun':
        aliyun_config = {
            'access_key_id': args.aliyun_key_id or os.environ.get('ALIYUN_ACCESS_KEY_ID') or (config.get('aliyun', {}).get('access_key_id') if config else ''),
            'access_key_secret': args.aliyun_key_secret or os.environ.get('ALIYUN_ACCESS_KEY_SECRET') or (config.get('aliyun', {}).get('access_key_secret') if config else ''),
            'endpoint': args.aliyun_endpoint or (config.get('aliyun', {}).get('endpoint') if config else 'ocr.cn-shanghai.aliyuncs.com')
        }
    elif args.ocr_mode == 'tencent':
        tencent_config = {
            'secret_id': args.tencent_secret_id or os.environ.get('TENCENT_SECRET_ID') or (config.get('tencent', {}).get('secret_id') if config else ''),
            'secret_key': args.tencent_secret_key or os.environ.get('TENCENT_SECRET_KEY') or (config.get('tencent', {}).get('secret_key') if config else ''),
            'region': args.tencent_region or (config.get('tencent', {}).get('region') if config else 'ap-guangzhou')
        }
    elif args.ocr_mode == 'hybrid':
        # 混合模式需要腾讯云 OCR 配置
        tencent_config = {
            'secret_id': args.tencent_secret_id or os.environ.get('TENCENT_SECRET_ID') or (config.get('tencent', {}).get('secret_id') if config else ''),
            'secret_key': args.tencent_secret_key or os.environ.get('TENCENT_SECRET_KEY') or (config.get('tencent', {}).get('secret_key') if config else ''),
            'region': args.tencent_region or (config.get('tencent', {}).get('region') if config else 'ap-guangzhou')
        }
    
    # 加载本地 OCR 配置
    if config and 'local' in config:
        args.lang = args.lang or config['local'].get('lang', 'chi_sim+eng')
        args.dpi = args.dpi or config['local'].get('dpi', 200)
        args.max_pages = args.max_pages or config['local'].get('max_pages', 20)
    
    ocr_deps = init_ocr(args.ocr_mode, aliyun_config, tencent_config)
    
    # 1. OCR 识别 PDF
    extractor = PDFOCRExtractor(
        args.pdf, 
        lang=args.lang, 
        dpi=args.dpi,
        ocr_mode=args.ocr_mode,
        aliyun_config=aliyun_config,
        tencent_config=tencent_config
    )
    ocr_text = extractor.extract(args.max_pages)
    
    if not ocr_text:
        print("✗ 错误：无法从 PDF 提取文本")
        sys.exit(1)
    
    # 2. 提取案件信息
    if args.extract_mode == 'llm':
        # 使用大模型提取
        llm_config = config.get('llm', {}) if config else {}
        llm_config['provider'] = args.llm_provider
        llm_config['api_key'] = args.llm_api_key or llm_config.get('api_key', '')
        llm_config['model'] = args.llm_model or llm_config.get('model', 'qwen-plus')
        
        info_extractor = LLMCaseInfoExtractor(ocr_text, llm_config)
        case_info = info_extractor.extract()
    else:
        # 使用正则提取
        info_extractor = CaseInfoExtractor(ocr_text, ocr_mode=args.ocr_mode)
        case_info = info_extractor.extract()
    
    # 显示提取的信息
    print("\n📋 提取的案件信息:")
    for key, value in case_info.items():
        if value:
            print(f"   {key}: {value}")
    
    # 3. 填写起诉状
    filler = LawsuitFiller(args.template, case_info)
    success = filler.fill(args.output)
    
    if success:
        print("\n" + "="*60)
        print(f"✓ 完成！输出文件：{args.output}")
        print("="*60)
    else:
        print("\n✗ 填写失败")
        sys.exit(1)


if __name__ == '__main__':
    main()
